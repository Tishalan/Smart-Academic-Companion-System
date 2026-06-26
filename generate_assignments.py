import os
import random
from datetime import datetime, timedelta
from app import create_app, db
from app.models import Assignment, Module, Student, Submission, Enrollment
from docx import Document

def generate_data():
    app = create_app('development')
    
    with app.app_context():
        # Ensure upload directories exist
        upload_dir = os.path.join('app', 'static', 'uploads')
        os.makedirs(upload_dir, exist_ok=True)
        
        # We will create an assignment for every module
        modules = Module.query.all()
        print(f"Found {len(modules)} modules. Generating assignments...")
        
        total_assignments = 0
        total_submissions = 0
        
        for idx, module in enumerate(modules):
            # Create dummy Word Document
            filename = f"Assignment_Brief_{module.module_code}.docx"
            filepath = os.path.join(upload_dir, filename)
            
            if not os.path.exists(filepath):
                doc = Document()
                doc.add_heading(f"Assignment Brief: {module.module_name}", 0)
                doc.add_paragraph(f"Module Code: {module.module_code}")
                doc.add_paragraph("Instructions:")
                doc.add_paragraph("1. Please complete all tasks outlined in this brief.")
                doc.add_paragraph("2. Ensure your submission is formatted correctly.")
                doc.add_paragraph("3. Upload your final document before the deadline.")
                doc.add_paragraph("\nGood luck!")
                doc.save(filepath)
            
            # Check if assignment exists for this module
            assignment = Assignment.query.filter_by(module_id=module.id).first()
            if not assignment:
                deadline = datetime.utcnow() + timedelta(days=random.randint(7, 30))
                assignment = Assignment(
                    module_id=module.id,
                    title=f"Final Assessment: {module.module_name}",
                    description="Please refer to the attached brief for detailed instructions.",
                    total_marks=100,
                    weightage=100.0,
                    deadline=deadline,
                    submission_type='file',
                    brief_file_path=filename
                )
                db.session.add(assignment)
                db.session.commit()
                total_assignments += 1
            elif not assignment.brief_file_path:
                assignment.brief_file_path = filename
                db.session.commit()
            
            # Generate a few dummy submissions from enrolled students so Lecturers have something to grade
            enrollments = Enrollment.query.filter_by(module_id=module.id).limit(5).all()
            for enrollment in enrollments:
                student = enrollment.student
                existing_sub = Submission.query.filter_by(assignment_id=assignment.id, student_id=student.id).first()
                if not existing_sub:
                    # Create a dummy submission doc
                    sub_filename = f"Sub_{student.student_id}_{module.module_code}.docx"
                    sub_filepath = os.path.join(upload_dir, sub_filename)
                    if not os.path.exists(sub_filepath):
                        sdoc = Document()
                        sdoc.add_heading(f"Submission for {module.module_name}", 0)
                        sdoc.add_paragraph(f"Student: {student.user.full_name} ({student.student_id})")
                        sdoc.add_paragraph("Here is my completed assignment.")
                        sdoc.save(sub_filepath)
                    
                    sub = Submission(
                        assignment_id=assignment.id,
                        student_id=student.id,
                        file_path=f"static/uploads/{sub_filename}",
                        submitted_at=datetime.utcnow() - timedelta(days=random.randint(1, 5)),
                        status='submitted'
                    )
                    db.session.add(sub)
                    total_submissions += 1
            
            db.session.commit()
            print(f"Processed module {idx+1}/{len(modules)}")
            
        print(f"Generated {total_assignments} assignments and {total_submissions} submissions.")

if __name__ == '__main__':
    generate_data()
